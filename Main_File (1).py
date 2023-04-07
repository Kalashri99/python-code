from selenium.webdriver.common.by import By
from selenium import webdriver
import pandas as pd
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
import docx
import re
import sys



def para_between_headings(file_path, heading_1, heading_2):
    doc = docx.Document(file_path)
    start_index, end_index = None, None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == heading_1:
            start_index = i
        elif para.text.strip() == heading_2:
            end_index = i
            break
    if start_index is not None and end_index is not None:
        return '\n'.join([para.text for para in doc.paragraphs[start_index+1:end_index]])
    else:
        return None


def data_between_headings(file_path, heading_1, heading_2):
    doc = docx.Document(file_path)
    start_index, end_index = None, None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == heading_1:
            start_index = i
        elif para.text.strip() == heading_2:
            end_index = i
            break
    if start_index is not None and end_index is not None:
        return '\n'.join([para.text for para in doc.paragraphs[start_index+1:end_index]])
    else:
        return None



def array_creator(result): 
    
    sentencearr = re.findall(r'\b.*?(?<=\n).*', result)  # array of all setences in testcase steps 
    textarr = re.findall(r'“(.*?)”', result)       # array of texts in para which are enclose with ""

    return sentencearr, textarr


def ExtractLocaters(sentencearr, textarr, data, worddict):  # data is dictionary of input and its value ,worddict->empty dictionary
    def GetExactXPath(element):
        xpath = driver.execute_script("function getElementXPath(elt) {var path = '';for (; elt && elt.nodeType == 1; elt = elt.parentNode) {index = getElementIndex(elt);var tagName = elt.nodeName.toLowerCase();var extra = '';if (index > 1) {extra = '[' + index + ']';}path = '/' + tagName + extra + path;}return path;}function getElementIndex(elt) {var index = 1;for (var sib = elt.previousSibling; sib; sib = sib.previousSibling) {if (sib.nodeType == 1 && sib.nodeName == elt.nodeName) {index++}}return index;}return getElementXPath(arguments[0]);", element)
        return xpath

    def label_function(word):
        try:
            labelPath = "//label[contains(text(),'{}')]".format(word)
            label = driver.find_element(By.XPATH, labelPath)
            inputId = label.get_attribute("for")
            inputPath = "input[id='{}']".format(inputId)
            ele = driver.find_element(By.CSS_SELECTOR, inputPath)
            return ele, "cssSelector", inputPath   # ele,locator type, locator value
        except:
            return None, None, None

    def text_function_textbox1(word):
        driver.implicitly_wait(10)
        # define a list of tags and attributes to search for
        input_tags = ["input","label","span","textarea"]
        input_attrs = ["placeholder", "name", "id"]
        # loop through the tags and attributes to find the input element
        for tag in input_tags:
            for attr in input_attrs:
                xpath = "//{}[@{}='{}'] | //{}[contains(normalize-space(),'{}')]/following::input[1]".format(tag, attr, word, tag, word, tag, attr, tag)
                try:
                    input_elem = driver.find_element(By.XPATH,xpath)
                    #break  # stop searching when the element is found
                    path=GetExactXPath(input_elem)
                    print(word +":"+path)
                    return input_elem,"xpath",xpath
                except:
                    continue  # try the next attribute or tag
        


    def text_function_textbox(word):
        try:
            try:
                driver.implicitly_wait(10)
                text="//input[@placeholder='{}']".format(word)
                ele = driver.find_element(By.XPATH, text)
                return ele, "xpath", text
            except:
                text = "// *[contains(text(),'{}')]".format(word)
                ele = driver.find_element(By.XPATH, text)
                if ele.tag_name == 'label':
                    return label_function(word)
                elif ele.tag_name == 'span' or ele.tag_name == 'div':
                    return span_function(ele)
        
        except:
            return None,None,None
    
    
    def text_function_button1(word):
        input_tags = ["input","button",'a',"label","span","textarea"]
        input_attrs = ["value", "name", "id"]
        # loop through the tags and attributes to find the input element
        for tag in input_tags:
            for attr in input_attrs:
                xpath = "//{}[@{}='{}'] | //{}[normalize-space()='{}']".format(tag, attr, word, tag, word)
                try:
                    input_elem = driver.find_element(By.XPATH,xpath)
                    #break  # stop searching when the element is found
                    path=GetExactXPath(input_elem)
                    print("path for "+word+":"+path)
                    return input_elem,"xpath",xpath
                except:
                    continue  # try the next attribute or tag
        return None,None,None

    def text_function_button(word):
        try:
            text="// *[contains(text(),'{}')]".format(word)
            elements=driver.find_elements(By.XPATH,text)
            for element in elements:
                tagname=element.tag_name.lower()
                if tagname in ['button','input','a','span','div']:
                    text="// {0}[contains(text(),'{1}')]".format(tagname,word)
                    return element,'xpath',text
        except:

           None,None,None
        '''
        try:
            text = "input[value='{}']".format(word)
            ele = driver.find_element(By.CSS_SELECTOR, text)
            return ele, "cssSelector", text
        except:
            try:
                text = "// a[contains(text(),'{}')]".format(word)
                ele = driver.find_element(By.XPATH, text)
                return ele, 'xpath', text
            except:
                try:
                    text = "// button[contains(text(),'{}')]".format(word)
                    ele = driver.find_element(By.XPATH, text)
                    return ele, 'xpath', text
                except:
                return "Error", "Error", "Error"
        '''



    def span_function(ele):
        try:
            parentEle = ele.find_element(By.XPATH, "..")
            inputId = parentEle.get_attribute("for")
            inputPath = "input[id='{}']".format(inputId)
            ele2 = driver.find_element(By.CSS_SELECTOR, inputPath)
            return ele2, "cssSelector", inputPath
        except:
            return None, None, None

    for i, j in zip(sentencearr, textarr):   # zip function to iterate 2 lists parallelly
        if 'URL' in i or '.com' in i or '.in' in i:
            '''
            chrome_options = Options()
            chrome_options.add_argument('--headless')
            driver = webdriver.Chrome(options=chrome_options)
            '''
            driver = webdriver.Chrome()  # driver to contol chrome browser
            
            driver.get(j)

            driver.execute_script("window.scrollTo(0,document.body.scrollHeight)")
            worddict = {i: {'Steps': i, 'Fields': "URL", 'Type': "URL",
                            'Data': j, 'Locater Type': '', 'Locater Value': ''}}
        elif 'Click' in i or 'Button' in i or 'click' in i or 'button' in i:
            ele, locater_type, locater_val = text_function_button1(j)
            worddict[i] = {'Steps': i, 'Fields': j, 'Type': "ClickJS", 'Data': '',
                           'Locater Type': locater_type, 'Locater Value': locater_val}
            driver.execute_script("window.scrollTo(0,document.body.scrollHeight)")
            ele.click()
        elif 'Enter' in i or 'Textbox' in i or 'enter' in i or 'textbox' in i:
            
            ele, locater_type, locater_val = text_function_textbox1(j)
            driver.execute_script("window.scrollTo(0,document.body.scrollHeight)")
            worddict[i] = {'Steps': i, 'Fields': j, 'Type': "Textbox", 'Data': data[j],
                           'Locater Type': locater_type, 'Locater Value': locater_val}
            ele.send_keys(data[j])
    return worddict  #dictionary of dicrionory


def toExcel(worddict,testName):
    df = pd.DataFrame.from_dict(worddict, orient='index')
    df.insert(0, 'TestCase Name', testName)
    writer = pd.ExcelWriter('D:\output.xlsx', engine='xlsxwriter')
    df.to_excel(writer, sheet_name='Sheet1', index=False)
    workbook = writer.book
    worksheet = writer.sheets['Sheet1']

    for i, col in enumerate(df.columns):
        column_len = df[col].astype(str).str.len().max()
        column_len = max(column_len, len(col))
        worksheet.set_column(i, i, column_len)
    writer.save()


file_path = "D:\RGI Project Documents\RGI-Feature testcase2.docx"
#file_path=sys.argv[1]

heading_1 = "TestCase Starts Here:"
heading_2 = "TestCase Ends Here:"
heading_1_data = "Inputs for above testcase starts here:"
heading_2_data = "Inputs for above testcase ends here:"
result = para_between_headings(file_path, heading_1, heading_2)
print(result)
Input = data_between_headings(file_path, heading_1_data, heading_2_data)
print(Input)
new_list = Input.split("\n")
split_list = [item.split(":") for item in new_list]
clean_list = [[item.replace('“', '') for item in inner_list]
              for inner_list in split_list]
final_list = [[item.replace('”', '') for item in inner_list]
              for inner_list in clean_list]
data = {item[0]: item[1] for item in final_list}




testName=result.split(":")[0]
sentencearr, textarr = array_creator(result)

allTestCases = {}
current_key = None




worddict = {}
worddict = ExtractLocaters(sentencearr, textarr, data, worddict)
toExcel(worddict,testName)

